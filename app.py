import os
import io
import re
import json
import time

import streamlit as st
import pandas as pd
import numpy as np
import altair as alt
import requests

from groq import Groq
from fastembed import TextEmbedding
from pypdf import PdfReader

st.set_page_config(page_title="DISCOM Data Chat", page_icon="⚡", layout="centered")
st.title("⚡ DISCOM Data Chat")
st.caption("Ask questions about the live feeder billing table, or upload documents to chat with them.")

# ========== CONFIG ==========
SUPABASE_URL = "https://hhwjxievppujzlnyqykp.supabase.co"
DB_TABLE_NAME = "feeder_billing_consumption"
DB_PAGE_SIZE = 1000

MODEL_PLANNER = "llama-3.1-8b-instant"     # structured JSON — small model is plenty
MODEL_ANSWER = "llama-3.3-70b-versatile"   # prose quality where it matters
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
TOP_K = 3
EMBED_BATCH_SIZE = 64
MAX_RESULT_ROWS = 200          # rows returned from a query
MAX_ROWS_TO_LLM = 15           # rows shown to the model for phrasing (token cost)
MAX_DISTINCT_LISTED = 25       # cardinality cutoff for listing values in the schema
MAX_HISTORY_MSGS = 2           # conversation turns replayed into the planner
MAX_HISTORY_CHARS = 200
MAX_CHART_ROWS = 40            # bars beyond this are unreadable

MONTHS = ["jan", "feb", "mar", "apr", "may", "jun",
          "jul", "aug", "sep", "oct", "nov", "dec"]
CHART_TYPES = {"bar", "line", "scatter", "area", "none"}

NUMERIC_COLS = [
    "mf", "initial_reading_kwh", "final_reading_kwh", "pf", "kwh", "kwh_exp",
    "kvah", "md_kw", "md_kva", "kvarh_q1", "kvarh_q2", "kvarh_q3", "kvarh_q4",
    "pwr_on_dur", "billi_year",
]

AGG_FUNCS = {"sum", "mean", "min", "max", "count", "nunique", "median", "std"}
OPS = {"==", "!=", ">", ">=", "<", "<=", "in", "not_in", "contains", "isnull", "notnull"}

TABLE_KNOWLEDGE = """
TABLE: feeder_billing_consumption — DISCOM feeder-level billing and consumption data

Column meanings:
- region, circle, division, zone: Administrative hierarchy of MP East DISCOM
- substation, substation_code, feeder, feeder_code: Physical infrastructure identifiers
- meter_serial_number: Unique meter ID on the feeder
- mf: Meter multiplication factor
- category: Feeder category (e.g., "In" = incoming feeder)
- bill_month, billi_year: Billing cycle month and year
- initial_reading_kwh, final_reading_kwh: Meter readings at start/end of billing cycle
- pf: Power factor (0 to 1); low PF indicates inefficient load
- kwh: Active energy consumed (kilowatt-hours) in the billing cycle
- kwh_exp: Exported energy (kWh) — relevant for feeders with reverse power flow
- kvah: Apparent energy (kilovolt-ampere-hours)
- md_kw, md_kva: Maximum demand recorded during the cycle (kW and kVA)
- kvarh_q1 to q4: Reactive energy by quadrant — used for power factor penalty calculations
- pwr_on_dur: Duration the feeder was powered on during the cycle
"""


def get_secret(name):
    try:
        if name in st.secrets:
            return st.secrets[name]
    except Exception:
        pass
    return os.getenv(name)


groq_key = get_secret("GROQ_API_KEY")
supabase_anon_key = get_secret("SUPABASE_ANON_KEY")

groq_client = None
if groq_key:
    try:
        groq_client = Groq(api_key=groq_key)
    except Exception as e:
        st.sidebar.error(f"Groq Error: {e}")


@st.cache_resource
def load_embedder():
    return TextEmbedding(model_name="BAAI/bge-small-en-v1.5")


# ========== SUPABASE → DATAFRAME ==========
def _supabase_headers(extra=None):
    headers = {
        "apikey": supabase_anon_key,
        "Authorization": f"Bearer {supabase_anon_key}",
    }
    if extra:
        headers.update(extra)
    return headers


@st.cache_data(ttl=3600, show_spinner=False, max_entries=1)
def load_dataframe():
    """Pull the whole table into memory once per hour. 8k rows is a few MB."""
    url = f"{SUPABASE_URL}/rest/v1/{DB_TABLE_NAME}?select=*"
    rows, offset = [], 0
    while True:
        headers = _supabase_headers({
            "Range-Unit": "items",
            "Range": f"{offset}-{offset + DB_PAGE_SIZE - 1}",
        })
        resp = requests.get(url, headers=headers, timeout=60)
        resp.raise_for_status()
        batch = resp.json()
        rows.extend(batch)
        if len(batch) < DB_PAGE_SIZE:
            break
        offset += DB_PAGE_SIZE

    df = pd.DataFrame(rows)
    for col in NUMERIC_COLS:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


@st.cache_data(ttl=3600, show_spinner=False, max_entries=1)
def build_schema_summary(_df, fingerprint):
    """
    Describe the table to the model as compactly as possible — this text is resent
    with every question, so every token here is paid for on each turn.
    """
    lines = [f"Table {DB_TABLE_NAME} ({len(_df)} rows). Columns:"]
    for col in _df.columns:
        s = _df[col]
        if pd.api.types.is_numeric_dtype(s):
            lines.append(f"{col} (num {s.min():.4g}..{s.max():.4g})")
        else:
            distinct = s.dropna().astype(str).str.strip().unique()
            if len(distinct) <= MAX_DISTINCT_LISTED:
                lines.append(f"{col} (text): {', '.join(sorted(distinct))}")
            else:
                lines.append(f"{col} (text, {len(distinct)} values)")
    return "\n".join(lines)


# ========== QUERY SPEC EXECUTION ==========
def _safe_name(name):
    return re.sub(r"\W+", "_", str(name)).strip("_") or "value"


def run_spec(df, spec, max_rows=MAX_RESULT_ROWS):
    """Execute a validated JSON query spec with pandas. No eval, no generated code."""
    out = df
    applied = []

    for f in spec.get("filters") or []:
        col, op, val = f.get("column"), f.get("op"), f.get("value")
        if col not in out.columns or op not in OPS:
            continue
        s = out[col]
        try:
            if op in {">", ">=", "<", "<="}:
                num = pd.to_numeric(s, errors="coerce")
                v = float(val)
                out = out[{">": num > v, ">=": num >= v, "<": num < v, "<=": num <= v}[op]]
            elif op == "==":
                out = out[s.astype(str).str.strip().str.lower() == str(val).strip().lower()]
            elif op == "!=":
                out = out[s.astype(str).str.strip().str.lower() != str(val).strip().lower()]
            elif op in {"in", "not_in"}:
                vals = [str(v).strip().lower() for v in (val if isinstance(val, list) else [val])]
                mask = s.astype(str).str.strip().str.lower().isin(vals)
                out = out[mask if op == "in" else ~mask]
            elif op == "contains":
                out = out[s.astype(str).str.contains(str(val), case=False, na=False)]
            elif op == "isnull":
                out = out[s.isna()]
            elif op == "notnull":
                out = out[s.notna()]
        except Exception:
            continue
        applied.append(f"{col} {op} {val}")

    matched_rows = len(out)

    group_by = [c for c in (spec.get("group_by") or []) if c in out.columns]
    aggs = [a for a in (spec.get("aggregations") or []) if a.get("func") in AGG_FUNCS]

    if aggs:
        named = {}
        for a in aggs:
            col = a.get("column")
            if col in out.columns:
                named[_safe_name(a.get("as") or f"{a['func']}_{col}")] = (col, a["func"])
        if named:
            if group_by:
                out = out.groupby(group_by, dropna=False).agg(**named).reset_index()
            else:
                out = pd.DataFrame({k: [getattr(out[v[0]], v[1])()] for k, v in named.items()})
    elif group_by:
        out = out.groupby(group_by, dropna=False).size().reset_index(name="row_count")
    else:
        cols = [c for c in (spec.get("columns") or []) if c in out.columns]
        if cols:
            out = out[cols]

    sort = spec.get("sort") or {}
    if sort.get("by") in out.columns:
        out = out.sort_values(sort["by"], ascending=not sort.get("desc", True), na_position="last")

    limit = spec.get("limit")
    limit = max_rows if limit is None else min(int(limit), max_rows)
    return out.head(limit), matched_rows, applied


# ========== CHARTING ==========
def month_order(values):
    """Return values ordered Jan..Dec if they all look like month names, else None."""
    vals = [str(v).strip() for v in values]
    keys = []
    for v in vals:
        k = v.lower()[:3]
        if k not in MONTHS:
            return None
        keys.append(MONTHS.index(k))
    return [v for _, v in sorted(zip(keys, vals))]


def infer_chart(result):
    """Fallback when the planner gives no usable chart: pick something sensible."""
    if result is None or len(result) < 2 or len(result) > 500:
        return None
    num = [c for c in result.columns if pd.api.types.is_numeric_dtype(result[c])]
    cat = [c for c in result.columns if c not in num]
    if cat and num:
        kind = "line" if month_order(result[cat[0]].unique()) else "bar"
        return {"type": kind, "x": cat[0], "y": num[0],
                "color": cat[1] if len(cat) > 1 else None}
    if len(num) >= 2:
        return {"type": "scatter", "x": num[0], "y": num[1], "color": None}
    return None


def build_chart(result, spec):
    """Turn a chart spec into an Altair chart. Returns (chart, note) or (None, None)."""
    if not spec or spec.get("type") in (None, "none"):
        return None, None
    kind, x, y, color = spec.get("type"), spec.get("x"), spec.get("y"), spec.get("color")
    if kind not in CHART_TYPES or x not in result.columns or y not in result.columns:
        return None, None
    if not pd.api.types.is_numeric_dtype(result[y]):
        return None, None

    data, note = result, None
    if len(data) > MAX_CHART_ROWS:
        data = data.head(MAX_CHART_ROWS)
        note = f"Chart shows the first {MAX_CHART_ROWS} of {len(result)} rows."

    x_is_num = pd.api.types.is_numeric_dtype(data[x])
    order = None if x_is_num else month_order(data[x].unique())

    if x_is_num:
        x_enc = alt.X(f"{x}:Q", title=x)
    elif order:
        x_enc = alt.X(f"{x}:N", sort=order, title=x)
    else:
        # preserve the order the query produced (usually sorted by the metric)
        x_enc = alt.X(f"{x}:N", sort=list(data[x].astype(str)), title=x)

    enc = {
        "x": x_enc,
        "y": alt.Y(f"{y}:Q", title=y),
        "tooltip": [alt.Tooltip(c) for c in data.columns[:6]],
    }
    if color and color in data.columns:
        enc["color"] = alt.Color(f"{color}:N", title=color)

    base = alt.Chart(data)
    marks = {
        "bar": base.mark_bar(),
        "line": base.mark_line(point=True),
        "area": base.mark_area(opacity=0.7),
        "scatter": base.mark_circle(size=70),
    }
    chart = marks[kind].encode(**enc).properties(
        height=340, title=spec.get("title") or ""
    ).interactive()
    return chart, note


def render_result(result, spec, key_prefix):
    """Render chart + table + download for a query result."""
    chart_spec = spec.get("chart") or infer_chart(result)
    chart, note = build_chart(result, chart_spec)
    if chart is None and chart_spec:
        # planner spec was unusable; try the inferred one before giving up
        chart, note = build_chart(result, infer_chart(result))
    if chart is not None:
        st.altair_chart(chart, use_container_width=True)
        if note:
            st.caption(note)
    with st.expander(f"Result table ({len(result)} rows)"):
        st.dataframe(result, use_container_width=True)
        st.download_button(
            "Download CSV",
            result.to_csv(index=False).encode("utf-8"),
            file_name="query_result.csv",
            mime="text/csv",
            key=f"dl_{key_prefix}",
        )


def call_llm_with_fallback(prompt, models, temperature=0.1, max_tokens=600):
    """
    Try each model in order. A 429 on one falls through to the next rather than
    failing the whole answer. Returns (text, model_used) or (None, None).
    """
    last = None
    for model in models:
        try:
            return call_llm(prompt, model=model, temperature=temperature,
                            max_tokens=max_tokens, retries=0), model
        except Exception as e:
            last = e
            if "429" in str(e) or "rate_limit" in str(e).lower():
                continue
            raise
    st.session_state.last_rate_error = str(last) if last else None
    return None, None


def summarize_result_offline(result, spec, matched, applied):
    """
    Build an answer from the result table with no API call at all.
    The numbers came from pandas, so this is exact — it just isn't prose.
    """
    parts = []
    restated = spec.get("restated")
    if restated:
        parts.append(f"**{restated}**")
    parts.append(
        f"{matched:,} rows matched"
        + (f" (filters: {'; '.join(applied)})" if applied else " (no filters)")
        + f", returning {len(result)} row(s)."
    )

    if len(result) == 1 and len(result.columns) <= 3:
        pairs = ", ".join(
            f"{c} = {result.iloc[0][c]:,.2f}" if pd.api.types.is_number(result.iloc[0][c])
            else f"{c} = {result.iloc[0][c]}"
            for c in result.columns
        )
        parts.append(f"Result: {pairs}")
    else:
        num_cols = [c for c in result.columns if pd.api.types.is_numeric_dtype(result[c])]
        cat_cols = [c for c in result.columns if c not in num_cols]
        if num_cols and cat_cols:
            metric, label = num_cols[0], cat_cols[0]
            top = result.head(3)
            lines = [f"- {r[label]}: {r[metric]:,.2f}" for _, r in top.iterrows()]
            parts.append(f"Top {len(lines)} by {metric}:\n" + "\n".join(lines))
            if len(num_cols) == 1:
                parts.append(f"Sum of {metric} across shown rows: {result[metric].sum():,.2f}")
    parts.append("_Chart and full table below._")
    return "\n\n".join(parts)


# ========== LLM CALLS ==========
def est_tokens(text):
    """Rough token estimate: ~4 chars per token. Good enough for budgeting."""
    return max(1, len(text) // 4)


def record_call(tokens):
    now = time.time()
    calls = [c for c in st.session_state.api_calls if now - c[0] < 60]
    calls.append((now, tokens))
    st.session_state.api_calls = calls


def usage_last_minute():
    now = time.time()
    recent = [c for c in st.session_state.api_calls if now - c[0] < 60]
    return len(recent), sum(t for _, t in recent)


def parse_retry_after(msg):
    for pattern in (r"try again in ([\d.]+)s", r"retry[-_ ]?after[\"':\s]+([\d.]+)"):
        m = re.search(pattern, msg, re.I)
        if m:
            return float(m.group(1))
    return None


def describe_rate_error(msg):
    """Groq's message says which limit tripped — surface it instead of guessing."""
    low = msg.lower()
    if "tokens per day" in low or "tpd" in low:
        kind = "daily token limit (TPD)"
    elif "tokens per minute" in low or "tpm" in low:
        kind = "tokens-per-minute limit (TPM)"
    elif "requests per day" in low or "rpd" in low:
        kind = "daily request limit (RPD)"
    elif "requests per minute" in low or "rpm" in low:
        kind = "requests-per-minute limit (RPM)"
    else:
        kind = "rate limit"
    wait = parse_retry_after(msg)
    suffix = f" Retry in about {wait:.0f}s." if wait else ""
    return f"Groq {kind} reached.{suffix}", msg


def call_llm(prompt, model, temperature=0.1, max_tokens=1024, retries=1):
    last_error = None
    for attempt in range(retries + 1):
        try:
            resp = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=model,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            text = resp.choices[0].message.content
            used = getattr(resp, "usage", None)
            record_call(used.total_tokens if used else est_tokens(prompt) + est_tokens(text))
            return text
        except Exception as e:
            last_error = e
            msg = str(e)
            if "429" not in msg and "rate_limit" not in msg.lower():
                raise
            wait = parse_retry_after(msg)
            if attempt < retries and wait is not None and wait <= 15:
                time.sleep(wait + 0.5)
                continue
            raise
    raise last_error


@st.cache_data(ttl=1800, show_spinner=False, max_entries=300)
def plan_query(question, history, schema, model):
    """Cached: asking the same question twice costs zero API calls."""
    raw = call_llm(PLANNER_TEMPLATE.format(
        knowledge=TABLE_KNOWLEDGE, schema=schema, history=history, question=question,
    ), model=model, temperature=0.0, max_tokens=450)
    return parse_json(raw) or {"route": "general"}


def parse_json(text):
    cleaned = re.sub(r"^```(?:json)?|```$", "", text.strip(), flags=re.MULTILINE).strip()
    match = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if not match:
        return None
    try:
        return json.loads(match.group(0))
    except json.JSONDecodeError:
        return None


PLANNER_TEMPLATE = """You convert a user's question into a JSON query specification for a pandas DataFrame.

{knowledge}

=== ACTUAL SCHEMA AND VALUES ===
{schema}

=== RECENT CONVERSATION ===
{history}

=== USER QUESTION ===
{question}

Decide the route:
- "data"    -> the question is about the table above (totals, counts, comparisons, rankings, lookups)
- "docs"    -> the question is about uploaded documents, not the table
- "general" -> neither; general knowledge

If route is "data", build the spec. Rules:
- Use ONLY column names that appear in the schema.
- Filter values MUST match the actual distinct values listed above, including capitalisation.
- Use aggregations for any question about totals, averages, counts, or rankings. Never return raw rows for those.
- func must be one of: sum, mean, min, max, count, nunique, median, std
- op must be one of: ==, !=, >, >=, <, <=, in, not_in, contains, isnull, notnull
- Set "limit" sensibly (e.g. 10 for "top 10", 50 otherwise).

Also choose a chart for the result:
- "bar"     -> comparing a metric across categories (division, region, circle), or a ranking
- "line"    -> a metric across bill_month or billi_year (a trend over time)
- "scatter" -> relationship between two numeric columns (e.g. pf vs kwh)
- "area"    -> cumulative or stacked totals over time
- "none"    -> the result is a single number, or a plain row lookup
Set "x" to the category or time column, "y" to the numeric metric (use the "as" alias
you defined in aggregations). Set "color" only when a second grouping column splits the
data into series; otherwise null.

Respond with JSON ONLY, no prose, no markdown fences:
{{
  "route": "data",
  "restated": "plain-English restatement of what will be computed",
  "filters": [{{"column": "bill_month", "op": "==", "value": "Mar"}}],
  "group_by": ["division"],
  "aggregations": [{{"column": "kwh", "func": "sum", "as": "total_kwh"}}],
  "columns": [],
  "sort": {{"by": "total_kwh", "desc": true}},
  "limit": 10,
  "chart": {{"type": "bar", "x": "division", "y": "total_kwh", "color": null,
             "title": "Total kWh by division, March 2026"}}
}}"""

ANSWER_TEMPLATE = """You are a power-distribution data analyst. Answer the user's question using the query result below.

=== WHAT WAS COMPUTED ===
{restated}
Filters applied: {applied}
Rows matching the filters: {matched}
Result rows shown: {shown}

=== QUERY RESULT (CSV) ===
{result}

=== USER QUESTION ===
{question}

Rules:
- The result table is authoritative. Do not invent, estimate, or extrapolate numbers not present in it.
- Report figures with units (kWh, kVA, etc.) and sensible rounding.
- If the result is empty, say so plainly and suggest which filter may be wrong.
- If the result was truncated, say the list is partial.
- Be concise. Lead with the answer, then supporting detail."""


# ========== DOCUMENT RAG (uploads only) ==========
def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    chunks, start = [], 0
    while start < len(text):
        end = min(start + size, len(text))
        piece = text[start:end].strip()
        if len(piece) > 50:
            chunks.append(piece)
        start = end - overlap if end < len(text) else end
    return chunks


def embed_chunks(embedder, chunks, show_progress=False, batch_size=EMBED_BATCH_SIZE):
    if not chunks:
        return [], []
    kept, vectors = [], []
    progress = st.progress(0) if show_progress else None
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start:start + batch_size]
        try:
            out = list(embedder.embed([c[:8000] for c in batch], batch_size=batch_size))
            kept.extend(batch)
            vectors.extend(np.asarray(v, dtype=np.float32) for v in out)
        except Exception as e:
            st.warning(f"Embed error near chunk {start}: {e}")
        if progress:
            progress.progress(min((start + batch_size) / len(chunks), 1.0))
    if progress:
        progress.empty()
    return kept, vectors


def add_docs(chunks, vectors, stat):
    if not vectors:
        return False
    matrix = np.array(vectors, dtype=np.float32)
    if st.session_state.doc_embeddings is None:
        st.session_state.doc_embeddings = matrix
        st.session_state.doc_chunks = list(chunks)
    else:
        st.session_state.doc_embeddings = np.vstack([st.session_state.doc_embeddings, matrix])
        st.session_state.doc_chunks.extend(chunks)
    st.session_state.file_stats.append(stat)
    return True


def retrieve_docs(embedder, query, top_k=TOP_K):
    matrix = st.session_state.doc_embeddings
    if matrix is None or not len(matrix):
        return [], 0.0
    q = np.asarray(list(embedder.embed([query[:8000]]))[0], dtype=np.float32)
    denom = np.linalg.norm(matrix, axis=1) * np.linalg.norm(q)
    denom[denom == 0] = 1e-9
    sims = (matrix @ q) / denom
    idx = np.argsort(sims)[-top_k:][::-1]
    return [(st.session_state.doc_chunks[i], float(sims[i])) for i in idx], float(sims[idx[0]])


# ========== EXTRACTORS ==========
def extract_pdf_text(file):
    reader = PdfReader(file)
    parts = []
    for i, page in enumerate(reader.pages):
        txt = page.extract_text()
        if txt and txt.strip():
            parts.append(f"[Page {i+1}]\n{txt}")
    return "\n\n".join(parts), len(reader.pages)


def extract_word_text(file):
    from docx import Document
    doc = Document(file)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        tables.append("\n".join(
            " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
        ))
    text = "\n\n".join(paragraphs)
    if tables:
        text += "\n\n[Tables]\n" + "\n\n".join(tables)
    return text, 1


def extract_excel_text(file):
    xls = pd.ExcelFile(file)
    parts = [
        f"[Sheet: {name}]\n{pd.read_excel(xls, sheet_name=name).to_string(index=False)}"
        for name in xls.sheet_names
    ]
    return "\n\n".join(parts), len(xls.sheet_names)


def extract_google_sheet(sheet_url):
    match = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", sheet_url)
    if not match:
        st.error("Invalid URL. Expected https://docs.google.com/spreadsheets/d/XXXX/edit")
        return None, 0
    export = f"https://docs.google.com/spreadsheets/d/{match.group(1)}/export?format=csv"
    resp = requests.get(export, timeout=30)
    resp.raise_for_status()
    return pd.read_csv(io.StringIO(resp.text)).to_string(index=False), 1


def process_file(file):
    name = file.name.lower()
    try:
        if name.endswith(".pdf"):
            return extract_pdf_text(file)
        if name.endswith(".docx"):
            return extract_word_text(file)
        if name.endswith((".xlsx", ".xls")):
            return extract_excel_text(file)
        if name.endswith(".csv"):
            return pd.read_csv(file).to_string(index=False), 1
        if name.endswith(".txt"):
            return file.read().decode("utf-8"), 1
    except Exception as e:
        st.error(f"Read error ({file.name}): {e}")
    return "", 0


# ========== SESSION STATE ==========
for key, default in [
    ("messages", []),
    ("doc_embeddings", None),
    ("doc_chunks", []),
    ("file_stats", []),
    ("api_calls", []),
    ("last_rate_error", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default


# ========== LOAD TABLE ==========
df, schema_summary, db_error = None, "", None
if supabase_anon_key:
    try:
        df = load_dataframe()
        schema_summary = build_schema_summary(df, fingerprint=(len(df), tuple(df.columns)))
    except Exception as e:
        db_error = str(e)


# ========== SIDEBAR ==========
with st.sidebar:
    st.header("Status")
    if groq_client:
        st.success("Groq connected")
    else:
        st.error("No Groq API key")

    if df is not None:
        st.success(f"Live table: {len(df):,} rows x {len(df.columns)} cols")
        if st.button("Refresh table", use_container_width=True):
            load_dataframe.clear()
            build_schema_summary.clear()
            st.rerun()
        with st.expander("Schema seen by the model"):
            st.code(schema_summary, language="text")
    elif db_error:
        st.error(f"DB error: {db_error}")
    else:
        st.warning("Supabase key not set")

    st.divider()
    st.subheader("Documents (optional)")
    st.caption("PDFs and reports use semantic search. The table above does not.")

    uploaded_files = st.file_uploader(
        "Upload files",
        type=["pdf", "docx", "xlsx", "xls", "csv", "txt"],
        accept_multiple_files=True,
        label_visibility="collapsed",
    )
    if uploaded_files and st.button("Index documents", use_container_width=True):
        embedder = load_embedder()
        with st.spinner("Indexing..."):
            for f in uploaded_files:
                text, pages = process_file(f)
                if not text:
                    continue
                kept, vectors = embed_chunks(embedder, chunk_text(text), show_progress=True)
                add_docs(kept, vectors, {"name": f.name, "pages": pages, "chunks": len(kept)})
        st.success(f"{len(st.session_state.doc_chunks)} chunks indexed")

    sheet_url = st.text_input("Google Sheet URL (public)", placeholder="https://docs.google.com/...")
    if st.button("Fetch sheet", use_container_width=True) and sheet_url:
        embedder = load_embedder()
        with st.spinner("Fetching..."):
            try:
                text, pages = extract_google_sheet(sheet_url)
                if text:
                    kept, vectors = embed_chunks(embedder, chunk_text(text), show_progress=True)
                    add_docs(kept, vectors, {"name": "Google_Sheet", "pages": pages, "chunks": len(kept)})
                    st.success("Sheet indexed")
            except Exception as e:
                st.error(f"Sheet error: {e}")

    if st.session_state.file_stats:
        st.divider()
        for s in st.session_state.file_stats:
            st.caption(f"{s['name']} - {s['chunks']} chunks")

    if st.session_state.messages:
        st.divider()
        if st.button("Clear chat", use_container_width=True):
            st.session_state.messages = []
            st.rerun()

    st.divider()
    st.subheader("API budget")
    calls_min, tokens_min = usage_last_minute()
    c1, c2 = st.columns(2)
    c1.metric("Calls (60s)", f"{calls_min}/30")
    c2.metric("Tokens (60s)", f"{tokens_min:,}")
    if tokens_min > 9000:
        st.warning("Approaching the per-minute token cap. Pause briefly.")
    st.caption(
        f"Planner: {MODEL_PLANNER}  \nAnswers: {MODEL_ANSWER}  \n"
        "Separate models draw on separate free-tier buckets. "
        "This meter counts this browser session only — Groq counts your whole account."
    )

    st.divider()
    narrate = st.checkbox(
        "Narrate answers with the LLM", value=True,
        help="Off = zero API calls for phrasing. Numbers, charts and tables still work fully.",
    )
    show_spec = st.checkbox(
        "Show query spec", value=False,
        help="Display the JSON the model generated for each question.",
    )


# ========== GUARDS ==========
if not groq_client:
    st.warning("Add GROQ_API_KEY and SUPABASE_ANON_KEY to your Streamlit secrets.")
    st.stop()

if df is None:
    st.warning("Could not load the database table. Check SUPABASE_ANON_KEY and table permissions.")
    st.stop()


# ========== CHAT ==========
for i, m in enumerate(st.session_state.messages):
    with st.chat_message(m["role"]):
        st.markdown(m["content"])
        if m.get("result_table"):
            past = pd.DataFrame(m["result_table"])
            render_result(past, {"chart": m.get("chart")}, key_prefix=f"hist_{i}")
        if m.get("spec"):
            with st.expander("Query spec"):
                st.code(json.dumps(m["spec"], indent=2), language="json")

if prompt := st.chat_input("e.g. total kWh by division for March 2026"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        try:
            history = "\n".join(
                f"{m['role']}: {m['content'][:MAX_HISTORY_CHARS]}"
                for m in st.session_state.messages[-(MAX_HISTORY_MSGS + 1):-1]
            )

            with st.spinner("Planning query..."):
                spec = plan_query(prompt, history, schema_summary, MODEL_PLANNER)

            route = spec.get("route", "general")
            record = {"role": "assistant", "spec": spec if show_spec else None,
                      "result_table": None, "chart": None}

            if route == "data":
                with st.spinner("Running query..."):
                    result, matched, applied = run_spec(df, spec)
                    shown = result.head(MAX_ROWS_TO_LLM)

                answer, model_used = None, None
                if narrate:
                    answer, model_used = call_llm_with_fallback(
                        ANSWER_TEMPLATE.format(
                            restated=spec.get("restated", "-"),
                            applied="; ".join(applied) or "none",
                            matched=f"{matched:,}",
                            shown=f"{len(shown)} of {len(result)}",
                            result=shown.to_csv(index=False) if len(shown) else "(empty)",
                            question=prompt,
                        ),
                        models=[MODEL_ANSWER, MODEL_PLANNER],
                        temperature=0.2, max_tokens=600,
                    )

                if answer is None:
                    answer = summarize_result_offline(result, spec, matched, applied)
                    st.markdown(answer)
                    if narrate:
                        st.warning(
                            "Both models are rate-limited, so this answer was built directly "
                            "from the query result. The figures are exact — only the wording "
                            "is unpolished."
                        )
                        err = st.session_state.get("last_rate_error")
                        if err:
                            with st.expander("Groq's exact response"):
                                st.code(err, language="text")
                else:
                    st.markdown(answer)
                    st.caption(f"Computed from {matched:,} matching rows · phrased by {model_used}")

                if len(result):
                    render_result(result, spec, key_prefix=f"live_{len(st.session_state.messages)}")
                    record["result_table"] = result.to_dict("records")
                    record["chart"] = spec.get("chart") or infer_chart(result)

            elif route == "docs" and st.session_state.doc_chunks:
                hits, best = retrieve_docs(load_embedder(), prompt)
                context = "\n\n---\n\n".join(c for c, _ in hits)
                answer, model_used = call_llm_with_fallback(
                    f"Answer using the document extracts below.\n\n=== EXTRACTS ===\n{context}\n\n"
                    f"=== QUESTION ===\n{prompt}\n\nIf the extracts do not answer it, say so.",
                    models=[MODEL_ANSWER, MODEL_PLANNER], max_tokens=600,
                )
                if answer is None:
                    raise RuntimeError(st.session_state.get("last_rate_error", "rate_limit_exceeded"))
                st.markdown(answer)
                st.caption(f"From documents (top score {best:.2f})")

            else:
                answer, model_used = call_llm_with_fallback(
                    f"=== RECENT CONVERSATION ===\n{history}\n\n=== QUESTION ===\n{prompt}\n\n"
                    "Answer clearly using general knowledge.",
                    models=[MODEL_ANSWER, MODEL_PLANNER], max_tokens=600,
                )
                if answer is None:
                    raise RuntimeError(st.session_state.get("last_rate_error", "rate_limit_exceeded"))
                st.markdown(answer)
                st.caption(f"General knowledge · {model_used}")

            if show_spec:
                with st.expander("Query spec"):
                    st.code(json.dumps(spec, indent=2), language="json")

            record["content"] = answer
            st.session_state.messages.append(record)

        except Exception as e:
            msg = str(e)
            if "429" in msg or "rate_limit" in msg.lower():
                headline, detail = describe_rate_error(msg)
                st.error(headline)
                with st.expander("Groq's exact response"):
                    st.code(detail, language="text")
                st.caption(
                    "Free-tier caps are per model. Adding a card to your Groq account "
                    "upgrades you to the Developer tier at no cost and raises limits substantially."
                )
            else:
                st.error(f"Error: {e}")
